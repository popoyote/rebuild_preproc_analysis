# for_pdfdoc.py
import pdfplumber
import PyPDF2  # Альтернатива pdfplumber
import docx2txt
from docx import Document  # Альтернатива для docx2txt


# --- PDF ---
def extract_text_pypdf2(file_path):
    """Извлечение текста через PyPDF2."""
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            return " ".join([page.extract_text() for page in reader.pages])
    except Exception as e:
        print(f"Ошибка PyPDF2: {e}")
        return ""


def extract_text_pdfplumber(file_path):
    """Извлечение текста через pdfplumber."""
    try:
        with pdfplumber.open(file_path) as pdf:
            return " ".join([page.extract_text() for page in pdf.pages])
    except Exception as e:
        print(f"Ошибка pdfplumber: {e}")
        return ""


# --- DOCX ---
def extract_text_docx(file_path):
    """Извлечение текста из DOCX с учетом таблиц."""
    doc = Document(file_path)
    text = []

    # Текст из параграфов
    for p in doc.paragraphs:
        text.append(" ".join(p.text.split()))

    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(" ".join(cell.text.split()))
    return "\n".join(text)


def extract_text_docx2txt(file_path):
    """Извлечение текста через docx2txt."""
    try:
        return docx2txt.process(file_path)
    except Exception as e:
        print(f"Ошибка docx2txt: {e}")
        return ""
